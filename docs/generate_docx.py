"""
Generate AgenticOps AI Word document (DOCX).
Run: python docs/generate_docx.py
Output: docs/AgenticOps_AI_Documentation.docx

Content synced from docs/doc_content.py (same source as PPTX).
"""

from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as exc:
    raise SystemExit("Install python-docx: pip install python-docx") from exc

from doc_content import PLATFORM_TITLE, PLATFORM_SUBTITLE, SECTIONS, GENERATED_DATE
from doc_outputs import DOCX_PATH, save_in_place

OUT = DOCX_PATH


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return h


def _add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def _add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True
    doc.add_paragraph()


def _build_document(dest: Path):
    doc = Document()

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(PLATFORM_TITLE)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(PLATFORM_SUBTITLE)
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Generated: {GENERATED_DATE}")
    meta_run.font.size = Pt(10)
    meta_run.italic = True

    doc.add_page_break()

    # Table of contents placeholder
    _add_heading(doc, "Table of Contents", level=1)
    for sec in SECTIONS:
        doc.add_paragraph(sec["title"], style="List Number")
    doc.add_page_break()

    # Sections
    for sec in SECTIONS:
        _add_heading(doc, sec["title"], level=1)
        for para in sec.get("paragraphs", []):
            p = doc.add_paragraph(para)
            for run in p.runs:
                run.font.size = Pt(11)
        if sec.get("table"):
            _add_table(doc, sec["table"])
        if sec.get("bullets"):
            _add_bullets(doc, sec["bullets"])
        doc.add_paragraph()

    # Footer note
    doc.add_page_break()
    _add_heading(doc, "Document Sync", level=1)
    note = doc.add_paragraph(
        "This document is auto-generated from docs/doc_content.py. "
        "To regenerate after feature changes, run: python docs/sync_all_documentation.py"
    )
    for run in note.runs:
        run.font.size = Pt(10)
        run.italic = True

    doc.save(str(dest))


def main():
    save_in_place(OUT, _build_document)
    print(f"[OK] Word document replaced in place: {OUT}")


if __name__ == "__main__":
    main()
